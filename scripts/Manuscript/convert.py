import os
import sys
import regex
import shutil
from docx import Document

DRAFT_DIR = '../../draft'
MANUSCRIPT_DIR = '../../Manuscript'

def find_simple_command(cmd, line):
    '''
    Finds commands in the form \<cmd_name>{<cmd_arg>}
    '''
    reg_exp = r'\\'+cmd+r'[ \t]*{(?P<arg>[^\\]*?)}'
    args, spans = [], []
    for match in regex.finditer(reg_exp, line):
        args.append(match.group('arg'))
        spans.append(match.span())
    return args, spans

def find_command(line):
    '''
    Finds commands in the form \<cmd_name>{<arg1>}{<arg2>}{<arg3>}
    '''
    reg_exp = r'\\(?P<cmd>[\w]+)[\s]*(({(?P<arg1>[^\\}{]*)})?)([\s]*({(?P<arg2>[^\\}{]*)})?)([\s]*({(?P<arg3>[^\\}{]*)})?)'
    matched_commands = []
    spans = []
    for match in regex.finditer(reg_exp, line, overlapped=False):
        matched_commands.append( (match.group('cmd'), match.group('arg1'), match.group('arg2'), match.group('arg3')) )
        spans.append(match.span())
    return matched_commands, spans

def find_param_command(cmd, line):
    '''
    Finds commands in the form \cmd[<arg>]{<arg1>}
    '''
    reg_exp = r'\\'+cmd+r'[\s]*(\[.*\])[\s]*(({(?P<arg1>[\w\W\/]*?)}))'
    matched_commands = []
    spans = []
    for match in regex.finditer(reg_exp, line, overlapped=False):
        matched_commands.append( match.group('arg1') )
        spans.append(match.span())
    return matched_commands, spans

def find_first_param_command(cmd, line):
    '''
    Finds commands in the form \cmd[<arg>]{<arg1>}
    '''
    reg_exp = r'\\'+cmd+r'[\s]*(\[(?P<opt>.*)\])[\s]*(({(?P<arg1>[\w\W\/]*?)}))'
    try:
        match = regex.finditer(reg_exp, line, overlapped=False).next()
    except:
        return None
    return match.group('arg1'), match.group('opt'), match.span()

def find_environment(env, string):
    '''
    Finds tex environments in the form:
    \begin{<env>}
    \end{<env}
    '''
    reg_exp = r'(\\begin[ \t]*{[ \t]*'+env+'[ \t]*}.*?\\end[ \t]*{[ \t]*'+env+'[ \t]*}?)'
    args, spans = [], []
    for match in regex.finditer(reg_exp, string, regex.DOTALL):
        spans.append(match.span())
    return spans

def remove_comments(string):
    '''
    Removes comments from string
    '''
    new_string = ''
    for line in string.split('\n'):
        comment_idx = line.find('%')
        if comment_idx!=-1:
            new_string += line[:comment_idx]+'\n'
        else:
            new_string += line+'\n'
    return new_string


def parse_table(tabular_env):
    tabular_regexp = r'\\begin\s*{\s*tabular\s*}\s*{(?P<arg>.*)}'
    match = regex.match(tabular_regexp, tabular_env)
    table_cols = regex.findall(r'[a-z]+', match.group('arg'))
    num_cols = len(table_cols)
    
    table_rows = regex.findall(r'(\\\\)+', tabular_env)
    num_rows = len(table_rows)

    #getting inside of the tabular environment
    reg_exp = r'\\begin\s*{\s*tabular\s*}{.*}(?P<inside>(.|\n|\r)*)\\end\s*{\s*tabular\s*}'
    match = regex.match(reg_exp, tabular_env)
    inside_tabular = match.group('inside')
    
    table = []
    
    for rown, line in enumerate(inside_tabular.split('\\\\')):
        table.append(['' for j in range(0,num_cols)])
        sline = line.replace('\\hline', '')
        sline = sline.split('&')
        for coln, subline in enumerate(sline):
            results = find_command(subline)
            for cmd, span in zip(results[0], results[1]):
                if cmd[0] == 'textbf':
                    subline = subline[:span[0]] + cmd[1] + subline[span[1]:]
                elif cmd[0] == 'multicolumn':
                    subline = subline[:span[0]] + cmd[3] + subline[span[1]:]
            table[-1][coln]+=subline.strip()
    
    return table
    
def resolve_inputs(input_file):
    '''
    Converts draft inputs to a single file
    '''
    resolved_file = ''
    for line in input_file:
        args, spans = find_simple_command('input', line)
        if len(args)>0: 
            for arg, span in zip(args, spans):
                filename = os.path.join(DRAFT_DIR, arg)
                insert = resolve_inputs(open(filename, 'r'))
                new_line = line[:span[0]] + insert + line[span[1]:]
        else:
            new_line = line
        resolved_file += new_line
    return resolved_file

def write_table_doc(name, table, filename):
    document = Document()
    document.add_heading(name)
    num_rows = len(table)
    num_cols = len(table[0])
    doc_table = document.add_table(rows=num_rows, cols=num_cols)
    for i in range(0,num_rows):
        for j in range(0,num_cols):
            cell = doc_table.cell(i, j)
            cell.text = table[i][j]
    document.save(filename)

def convert_tables(input_string):
    for n, span in enumerate(find_environment('table', input_string)):
        table_env = input_string[span[0]:span[1]]
        table_env = remove_comments(table_env)
        tabular_span = find_environment('tabular', table_env)
        tabular_env = table_env[tabular_span[0][0]:tabular_span[0][1]]
        table = parse_table(tabular_env)
        name = 'Table%d'%(n+1)
        write_table_doc(name, table, os.path.join(MANUSCRIPT_DIR, name+'.doc'))

def check_image(filename):
    # if filename.find('.eps')==-1:
    #     print 'Wrong format'
    #     return False
    
    print 'Image is fine'
    return True

def find_images(input_string):
    image_index = 1
    output_string = ''
    start_index = 0
    func = find_first_param_command('includegraphics', input_string[start_index:])
    while( not (func is None) ):
        span = func[2]

        image_rel_path = func[0]
        old_path = os.path.join(DRAFT_DIR, image_rel_path)
        new_rel_path = os.path.join('image%d.eps'%(image_index))
        new_path = os.path.join(MANUSCRIPT_DIR, new_rel_path)
        print 'Image: %s'%old_path, '->', new_path
        shutil.copy(old_path, new_path)

        output_string += input_string[start_index:start_index+span[0]]
        output_string += '\\includegraphics[%s]{%s}'%(func[1], new_rel_path)
        start_index += span[1]
        func = find_first_param_command('includegraphics', input_string[start_index:])
        image_index+=1
    
    output_string += input_string[start_index:]
    
    return output_string

if __name__ == '__main__':
    main_draft_path = os.path.join(DRAFT_DIR, 'main.tex')
    main_man_path = os.path.join(MANUSCRIPT_DIR, 'MainDocument.tex')

    # Put every input into single document
    manuscript = resolve_inputs(open(main_draft_path, 'r'))
    
    # Copy bibliography
    match = find_simple_command('bibliography', manuscript)
    bib_file = match[0][0]
    shutil.copy(os.path.join(DRAFT_DIR, bib_file), os.path.join(MANUSCRIPT_DIR, bib_file))

    # Remove comments
    # manuscript = remove_comments(manuscript)
    # Convert tables to *.doc
    convert_tables(manuscript)
    # Convert images
    manuscript = find_images(manuscript)
    
    
    
    with open(main_man_path, 'w') as fout:
        fout.write(manuscript)

    